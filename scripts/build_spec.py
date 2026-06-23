# -*- coding: utf-8 -*-
"""Render Bharat Direct Open Protocol spec + README to .md and .docx from one source."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\dhire\OneDrive\Documents\00 Dhirendra\Innovations\Payment Innovation\Bharat Direct'
NAVY = RGBColor(0x11,0x18,0x3A); BLUE = RGBColor(0x2F,0x4B,0x8F)
AMBER = RGBColor(0xB0,0x6A,0x00); GREY = RGBColor(0x55,0x5A,0x66); WHITE = RGBColor(0xFF,0xFF,0xFF)

# ---------- block helpers ----------
def H1(t): return ('h1',t)
def H2(t): return ('h2',t)
def H3(t): return ('h3',t)
def P(t):  return ('p',t)
def BUL(items): return ('bul',items)     # items: list of str OR (lead, rest)
def NUM(items): return ('num',items)
def CODE(t): return ('code',t.strip('\n'))
def TBL(h,rows): return ('table',h,rows)
def RULE(): return ('rule',)

# ===================== SPEC CONTENT =====================
SPEC = [
H1('Bharat Direct — Open Protocol Specification'),
P('Version 0.1 (Working Draft). An open standard for aggregator-free, bank-to-bank merchant payments with risk-aware authorization and programmable split settlement. This document is a discussion draft for the payments ecosystem; it is not a product and not final.'),
RULE(),

H2('1. Purpose and design goals'),
P('Bharat Direct defines how a merchant accepts a payment directly through a bank (no payment aggregator in the money path), how the rail attaches a fraud risk signal at the moment of authorization, and how a single customer payment can settle to multiple payees. It is designed to sit on top of existing real-time rails (UPI-class), reusing their authentication, mandates and tokenization rather than replacing them.'),
BUL([
 ('No rent layer. ','A merchant connects to a sponsor bank / PSP directly. No aggregator collects a percentage in between.'),
 ('Risk-aware by default. ','Every authorization can carry a risk score that drives on-screen interventions (the Aware Layer).'),
 ('Programmable settlement. ','One payment may split natively to many payees in a single netted cycle, removing the need for pooled escrow.'),
 ('Reuse, don’t rebuild. ','Authentication, mandates, tokenization and identity (CKYC) are referenced, not reinvented.'),
 ('Spec first. ','Conformance is defined by this document so any bank, merchant platform or app can interoperate.'),
]),

H2('2. Actors'),
BUL([
 ('Payer. ','The customer paying, holding an account at a Payer Bank.'),
 ('Payee(s). ','One or more recipients (merchant, marketplace sellers, commission account).'),
 ('Merchant / Platform. ','Initiates the payment request; may be a single seller or a marketplace.'),
 ('Sponsor Bank / PSP. ','The consortium bank that gives the merchant rail access. Replaces the PA.'),
 ('Switch. ','Thin, neutral router and net-settlement engine. Never custodies funds.'),
 ('FIC. ','Fraud Intelligence Consortium — returns a risk score; never receives raw PII.'),
 ('Aware Layer. ','Client-side component in the payer app that renders interventions based on the score.'),
]),

H2('3. Architecture overview'),
CODE('''Merchant/Platform                 Sponsor Bank / Switch              Payer App
      |                                    |                              |
      |  POST /payments (+split manifest)  |                              |
      |----------------------------------->|                              |
      |                                    |  collect / intent --------->  |
      |                                    |        POST /risk/score ----> FIC
      |                                    |  <---------- riskScore ------- |
      |                                    |  authorize  [Aware Layer UI]  |
      |                                    |  <---- payerAuth (PIN/bio) --- |
      |                                    |  debit -> net settle -> credit(N payees)
      |  <---------- webhook: SETTLED -----|                              |'''),

H2('4. Core resources and endpoints'),
P('All endpoints are HTTPS, JSON, and require mutual authentication and request signing (Section 9). Base path: /v1.'),
TBL(['Method & path','Purpose'],[
 ['POST /payments','Initiate a payment (single or split). Returns a payment with state INITIATED.'],
 ['GET /payments/{id}','Fetch payment state and settlement legs.'],
 ['POST /payments/{id}/refund','Reverse all or part of a settled payment.'],
 ['POST /risk/score','FIC: return a risk band + reasons for a proposed authorization.'],
 ['POST /mandates','Create a recurring mandate (AutoPay-class).'],
 ['POST /webhooks/test','Verify a merchant/bank webhook endpoint.'],
]),

H2('5. The payment request (with split manifest)'),
P('A marketplace expresses the split inline. The sum of split amounts MUST equal amount. If splits is omitted, the whole amount settles to a single default payee.'),
CODE('''POST /v1/payments
{
  "merchantId": "amzn-in",
  "orderRef": "AMZ-7781",
  "amount": 1000,
  "currency": "INR",
  "payer": { "handle": "customer@psp" },           // or null for intent flow
  "capture": "AUTO",
  "splits": [
    { "payee": "sellerA@bank", "amount": 400, "purpose": "GOODS" },
    { "payee": "sellerB@bank", "amount": 450, "purpose": "GOODS" },
    { "payee": "amzn-comm@bank", "amount": 150, "purpose": "COMMISSION" }
  ],
  "callbackUrl": "https://merchant.example/bd/webhook",
  "idempotencyKey": "AMZ-7781-1"
}'''),
P('Response:'),
CODE('''201 Created
{
  "paymentId": "pay_01HF...",
  "state": "INITIATED",
  "authUrl": "upi://pay?...",        // intent deep link, when payer not pre-bound
  "expiresAt": "2026-06-23T10:05:00Z"
}'''),

H2('6. The risk score (FIC) and the Aware Layer'),
P('Before the payer authorizes, the Payer Bank calls the FIC. The FIC returns a band and machine-readable reasons; it does NOT return raw data about other customers. The Aware Layer maps the band to an intervention.'),
CODE('''POST /v1/risk/score
{
  "payerRef": "hash:9af1...",        // pseudonymous
  "payeeRef": "sellerA@bank",
  "amount": 1000,
  "context": { "newPayee": true, "channel": "INTENT", "screenShareDetected": false }
}
--> 200
{
  "band": "HIGH",                    // LOW | ELEVATED | HIGH | BLOCK
  "score": 0.86,
  "reasons": ["PAYEE_RECENT_FRAUD_REPORTS", "NEW_PAYEE", "AMOUNT_ANOMALY"],
  "payeeReports24h": 9
}'''),
P('Intervention mapping (a conforming Aware Layer MUST implement at least these):'),
TBL(['Band','Required client behaviour'],[
 ['LOW','Frictionless. No interruption.'],
 ['ELEVATED','Named-threat notice + payee reputation shown; single confirm.'],
 ['HIGH','Unskippable cooling countdown + trance-breaker question ("Is someone instructing you to pay right now?") + active-recall confirm.'],
 ['BLOCK','Prevent authorization; route to helpline / cooling period; log to FIC.'],
]),
P('Intervention outcome is reported back so the FIC learns which interventions stopped which losses:'),
CODE('''POST /v1/risk/outcome
{ "scoreId": "rsk_01...", "shown": "HIGH", "userAction": "ABORTED", "abortReason": "TRANCE_BREAKER" }'''),

H2('7. Settlement model'),
P('The Switch computes net positions across all banks for a cycle and settles directly to each payee account. There is no pooled escrow and no intermediary float. Each split becomes an independent settlement leg with its own state.'),
CODE('''GET /v1/payments/pay_01HF...
{
  "paymentId": "pay_01HF...",
  "state": "SETTLED",
  "legs": [
    { "payee": "sellerA@bank", "amount": 400, "state": "SETTLED", "utr": "..." },
    { "payee": "sellerB@bank", "amount": 450, "state": "SETTLED", "utr": "..." },
    { "payee": "amzn-comm@bank", "amount": 150, "state": "SETTLED", "utr": "..." }
  ]
}'''),

H2('8. Refunds, reversals and disputes'),
BUL([
 ('Refund. ','POST /payments/{id}/refund with optional per-leg amounts; issued as a reversal message on the rail, not a card-style chargeback.'),
 ('Dispute. ','Raised through the consortium rulebook and online dispute resolution (ODR); the protocol carries dispute state but arbitration is governed off-band.'),
 ('Liability. ','For APP fraud, the protocol records the receiving-bank reference so a shared-liability rulebook can apportion loss (see whitepaper, Section 5).'),
]),

H2('9. Security and privacy'),
BUL([
 ('Transport & auth. ','Mutual TLS; every request signed (detached JWS). Replay-protected with idempotencyKey and timestamp.'),
 ('Idempotency. ','POST /payments and refunds MUST be idempotent on idempotencyKey.'),
 ('PII minimization. ','The FIC receives pseudonymous references (hashed), never raw identity. Identity resolution stays with CKYC and the Payer Bank.'),
 ('Federated learning. ','Member banks contribute model gradients/signals, not raw customer rows. Conformant FIC nodes MUST NOT export raw PII across members.'),
 ('Auditability. ','Every score, intervention shown, and outcome is logged for supervision.'),
]),

H2('10. Error model'),
CODE('''4xx/5xx
{ "error": "SPLIT_SUM_MISMATCH", "message": "splits must equal amount", "traceId": "..." }'''),
TBL(['Code','Meaning'],[
 ['SPLIT_SUM_MISMATCH','Sum of splits != amount.'],
 ['PAYEE_BLOCKED','A payee is BLOCK-banded by the FIC.'],
 ['IDEMPOTENCY_CONFLICT','Reused key with different body.'],
 ['MANDATE_REQUIRED','Recurring debit without a valid mandate.'],
]),

H2('11. Conformance profiles'),
P('An implementation declares one or more profiles so partners know what it supports:'),
BUL([
 ('Profile M (Merchant). ','Can initiate single-payee payments and handle webhooks.'),
 ('Profile S (Split). ','Adds programmable split settlement.'),
 ('Profile A (Aware). ','Implements the risk call and the full intervention mapping in Section 6.'),
 ('Profile F (FIC node). ','Serves /risk/score and participates in federated learning without exporting PII.'),
]),

H2('12. Reference implementation (planned modules)'),
TBL(['Module','Role','OSS'],[
 ['bharat-direct-merchant','Server SDK + e-commerce plugins (WooCommerce/Shopify/Magento) — Profile M/S','Yes'],
 ['bharat-direct-aware','Drop-in Aware Layer UI + risk client — Profile A','Yes'],
 ['bharat-direct-fic','Risk-score API + mule-graph schema + federated-learning reference — Profile F (mock/sandbox first)','Yes (reference)'],
 ['bharat-direct-spec','This document + OpenAPI + JSON schemas + conformance tests','Yes'],
]),

H2('13. Versioning and governance'),
P('Semantic versioning on the protocol. Breaking changes increment the major version and the /v{n} base path. The specification is intended to be stewarded by a neutral, not-for-profit consortium (NPCI-style governance) with open membership, so no single vendor controls it. Until then this is an open discussion draft released for comment.'),
RULE(),
P('See the accompanying white paper "Bharat Direct — Finishing what UPI started" for the economic and policy rationale, sourced figures, and the fraud-psychology basis for the Aware Layer.'),
]

# ===================== README CONTENT =====================
README = [
H1('Bharat Direct'),
P('An open protocol and reference toolkit for aggregator-free, bank-to-bank merchant payments — with a fraud-intelligence layer that warns users at the one moment that matters: the instant they authorize a payment.'),
P('Status: working draft / discussion starter. Not a product. Released for the ecosystem to examine, challenge, and build on.'),
RULE(),

H2('Why'),
P('India built the hardest 80% of a modern payments system — UPI, a bank-owned NPCI, and central CKYC. Two inefficiencies remain: a payment-aggregator layer that still collects rent the network no longer needs, and authorized push-payment (APP) fraud that no faster rail can stop. Bharat Direct addresses both.'),
BUL([
 ('Pillar 1 — Direct Settlement. ','Merchants connect to banks directly; funds settle bank-to-bank; only identity stays central. No PA money-path, no escrow.'),
 ('Pillar 2 — Fraud Intelligence Consortium (FIC). ','A shared, privacy-preserving risk signal that powers on-screen interventions (the Aware Layer) at the moment of authorization.'),
]),

H2('Architecture'),
CODE('''            +---------------------- Aware Layer ----------------------+
            |   dynamic friction on the payment screen (risk-driven)  |
            +----------------^------------------------^---------------+
                             | risk score             | pay / settle
            +----------------+--------+      +---------+----------------+
            |  FIC (Pillar 2)         |      |  Direct Settlement (P1)  |
            |  mule graph, typologies |      |  thin switch, net settle |
            |  federated ML, scoring  |      |  CKYC central only       |
            +-------------------------+      +--------------------------+
                          \\___ co-owned, NPCI-style governance ___/'''),

H2('Repository layout'),
CODE('''bharat-direct/
  spec/            # the open protocol (SPEC.md, OpenAPI, JSON schemas, conformance tests)
  merchant/        # server SDK + e-commerce plugins (Profile M/S) -- accept pay-by-bank, no PA
  aware/           # Aware Layer SDK: intervention UI + risk client (Profile A)
  fic/             # Fraud Intelligence reference service + mock sandbox (Profile F)
  examples/        # demo merchant + demo payer app wired end-to-end against the mock FIC
  docs/            # whitepaper, deck, this README''')
,

H2('What you can build today (no permission needed)'),
NUM([
 'The merchant connector against a bank/PSP UPI sandbox — proves merchants onboard without a PA.',
 'A demo payer app with the Aware Layer interventions firing on a mock risk score — demonstrable from a laptop.',
 'The published spec + conformance tests — the artifact others adopt without you.',
]),
P('What needs the ecosystem (not blocking the MVP): native split settlement (an NPCI rail feature) and a real FIC with live bank data sharing (an institutional change). Build the demonstrable pieces first; use the evidence to earn the rest.'),

H2('Quick start (reference, planned)'),
CODE('''# run the mock FIC + demo apps
git clone https://github.com/<org>/bharat-direct
cd bharat-direct
make demo        # starts mock FIC, a demo merchant, and a demo payer app
# open the payer app, pay a "high-risk" test payee, watch the Aware Layer intervene'''),

H2('Conformance profiles'),
TBL(['Profile','Capability'],[
 ['M','Merchant: initiate single-payee payments + webhooks'],
 ['S','Split: programmable multi-payee settlement'],
 ['A','Aware: risk call + full intervention mapping'],
 ['F','FIC node: serve risk scores, federated learning, no PII export'],
]),

H2('Governance and license'),
BUL([
 ('Spec. ','Intended for a neutral, not-for-profit, open-membership steward (NPCI-style). Suggested license for the specification: CC BY 4.0.'),
 ('Code. ','Suggested license for reference implementations: Apache-2.0.'),
 ('Principle. ','This is meant to become shared public infrastructure. Contributions and critique are the point.'),
]),

H2('Contributing'),
P('Open an issue describing the flow or risk you want to change. Protocol changes go through the spec/ directory with conformance tests. Be specific: cite the section, the actor, and the failure mode. The fraud-reduction claim must be measurable — proposals that affect the Aware Layer should describe how the effect would be measured.'),
RULE(),
P('Read the white paper for the full rationale, sourced figures, and the psychology behind the Aware Layer.'),
]

# ---------- markdown renderer ----------
def to_md(blocks):
    out=[]
    for b in blocks:
        k=b[0]
        if k=='h1': out.append(f'# {b[1]}\n')
        elif k=='h2': out.append(f'## {b[1]}\n')
        elif k=='h3': out.append(f'### {b[1]}\n')
        elif k=='p': out.append(f'{b[1]}\n')
        elif k=='code': out.append('```\n'+b[1]+'\n```\n')
        elif k=='rule': out.append('---\n')
        elif k=='bul':
            for it in b[1]:
                if isinstance(it,tuple): out.append(f'- **{it[0].rstrip()}** {it[1].lstrip()}')
                else: out.append(f'- {it}')
            out.append('')
        elif k=='num':
            for i,it in enumerate(b[1]): out.append(f'{i+1}. {it}')
            out.append('')
        elif k=='table':
            hdr=b[1]; rows=b[2]
            out.append('| '+' | '.join(hdr)+' |')
            out.append('| '+' | '.join(['---']*len(hdr))+' |')
            for r in rows: out.append('| '+' | '.join(r)+' |')
            out.append('')
    return '\n'.join(out)+'\n'

# ---------- docx renderer ----------
def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def to_docx(blocks, path, title):
    doc=Document(); n=doc.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(10.5)
    n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.12
    def runp(text,size=10.5,bold=False,italic=False,color=None):
        p=doc.add_paragraph(); r=p.add_run(text); f=r.font
        f.size=Pt(size); f.bold=bold; f.italic=italic
        if color: f.color.rgb=color
        return p
    for b in blocks:
        k=b[0]
        if k=='h1': p=runp(b[1],20,True,color=NAVY); p.paragraph_format.space_after=Pt(8)
        elif k=='h2': p=runp(b[1],14.5,True,color=BLUE); p.paragraph_format.space_before=Pt(12)
        elif k=='h3': p=runp(b[1],12,True,color=NAVY); p.paragraph_format.space_before=Pt(8)
        elif k=='p': runp(b[1])
        elif k=='rule':
            p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr')
            bo=OxmlElement('w:bottom'); bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'6')
            bo.set(qn('w:space'),'1'); bo.set(qn('w:color'),'2F4B8F'); pb.append(bo); pPr.append(pb)
        elif k=='bul':
            for it in b[1]:
                p=doc.add_paragraph(style='List Bullet')
                if isinstance(it,tuple):
                    r=p.add_run(it[0]); r.bold=True; p.add_run(it[1])
                else: p.add_run(it)
        elif k=='num':
            for it in b[1]:
                doc.add_paragraph(it, style='List Number')
        elif k=='code':
            t=doc.add_table(rows=1,cols=1); t.style='Table Grid'
            cell=t.rows[0].cells[0]; shade(cell,'F2F4F8'); cell.text=''
            for i,ln in enumerate(b[1].split('\n')):
                para=cell.paragraphs[0] if i==0 else cell.add_paragraph()
                para.paragraph_format.space_after=Pt(0); para.paragraph_format.line_spacing=1.0
                r=para.add_run(ln if ln else ' '); r.font.name='Consolas'; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x1A,0x1F,0x2E)
            doc.add_paragraph()
        elif k=='table':
            hdr=b[1]; rows=b[2]; t=doc.add_table(rows=1,cols=len(hdr)); t.style='Table Grid'
            for i,htext in enumerate(hdr):
                c=t.rows[0].cells[i]; shade(c,'11183A'); c.text=''
                rr=c.paragraphs[0].add_run(htext); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=WHITE
            for row in rows:
                cells=t.add_row().cells
                for i,val in enumerate(row):
                    cells[i].text=''; rr=cells[i].paragraphs[0].add_run(val); rr.font.size=Pt(9)
            doc.add_paragraph()
    doc.save(path)

# ---------- write all 4 files ----------
with open(os.path.join(OUT,'SPEC.md'),'w',encoding='utf-8') as f: f.write(to_md(SPEC))
with open(os.path.join(OUT,'README.md'),'w',encoding='utf-8') as f: f.write(to_md(README))
to_docx(SPEC, os.path.join(OUT,'Bharat-Direct-Spec.docx'), 'Spec')
to_docx(README, os.path.join(OUT,'Bharat-Direct-README.docx'), 'README')
print('Wrote: SPEC.md, README.md, Bharat-Direct-Spec.docx, Bharat-Direct-README.docx')
