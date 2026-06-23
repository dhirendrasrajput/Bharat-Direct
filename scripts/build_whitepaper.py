# -*- coding: utf-8 -*-
"""Bharat Direct white paper (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x11, 0x18, 0x3A)
BLUE = RGBColor(0x2F, 0x4B, 0x8F)
AMBER = RGBColor(0xC8, 0x7A, 0x00)
GREY = RGBColor(0x55, 0x5A, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
# base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, size=10, white=False, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
    if white: r.font.color.rgb = WHITE
    elif color: r.font.color.rgb = color

def heading(text, size=15, color=NAVY, space_before=14, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    r.font.name = 'Calibri'
    return p

def sub(text, size=12, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def body(text, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def rule_line():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '2F4B8F')
    pbdr.append(bottom); pPr.append(pbdr)

# ---------------- TITLE PAGE ----------------
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BHARAT DIRECT'); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Finishing what UPI started'); r.font.size = Pt(16); r.italic = True; r.font.color.rgb = AMBER
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A bank-owned, aggregator-free payments rail —\nand a Fraud Intelligence Consortium to close the human gap UPI left open')
r.font.size = Pt(13); r.font.color.rgb = GREY
for _ in range(8): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A discussion paper for regulators, banks and the payments industry'); r.font.size = Pt(11); r.font.color.rgb = GREY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Working draft · For circulation and comment'); r.font.size = Pt(10); r.italic = True; r.font.color.rgb = GREY
doc.add_page_break()

# ---------------- EXECUTIVE SUMMARY ----------------
heading('Executive summary', 18)
rule_line()
body('India has built the hardest 80% of a modern, low-cost payments system: a real-time rail (UPI) processing over 21 billion transactions a month, a bank-owned not-for-profit operator (NPCI), and a central identity registry (CKYC/CERSAI). Two structural inefficiencies remain.')
bullet(' a private payment-aggregator (PA) layer still collects intermediation rent between merchants and banks, even though the money itself can move bank-to-bank.', bold_lead='The toll-booth:')
bullet(' authorized push-payment (APP) / social-engineering fraud — where a victim is psychologically manipulated into authorising a real payment — is the fastest-growing financial crime, and the one gap a faster rail does nothing to close.', bold_lead='The fraud gap:')
body('Bharat Direct proposes two co-owned, NPCI-style utilities that solve both, and which only work at network scale:')
bullet(' merchants connect directly to banks; funds settle bank-to-bank over a thin switch; only identity (CKYC) stays central. The PA money-path — and its escrow regime — disappears.', bold_lead='Pillar 1 — Direct Settlement Consortium (DSC):')
bullet(' a shared, privacy-preserving fraud utility — a real-time mule/payee reputation graph, a scam-typology library, and federated machine learning — that powers an on-screen "Aware Layer" delivering precise, in-the-moment warnings at the exact instant of payment authorisation.', bold_lead='Pillar 2 — Fraud Intelligence Consortium (FIC):')
body('Together they answer the only serious objection to removing aggregators — "cheaper, but who handles fraud?" — and they handle it better than any single intermediary could, because fraud is a network phenomenon that no isolated institution can see whole.')

# ---------------- THE PROBLEM ----------------
heading('1.  The problem: rent and fear', 16)
rule_line()
sub('Where the cost actually sits')
body('On a typical card transaction, the merchant fee splits roughly into interchange to the issuing bank (60–80%), a scheme fee to the card network (5–15%), and the acquirer-plus-aggregator margin (10–25%). Only the last of these is pure intermediation rent — the interchange funds genuine credit risk; the scheme fee funds a network that a domestic rail can replace. For account-based payments, UPI already proves the rent can fall to near-zero.')
sub('Why fraud is the harder half')
body('A faster, cheaper rail does not stop a human being from being talked into paying a criminal. APP fraud now dominates real-time-payment fraud worldwide. It cannot be solved by moving money faster — only by intervening at the moment a person authorises the payment.')

# ---------------- THESIS ----------------
heading('2.  The thesis: complete what India already built', 16)
rule_line()
body('Bharat Direct is not a new paradigm. NPCI is already a bank consortium. CERSAI is already a central identity utility. UPI already moves money bank-to-bank with no aggregator in the path for person-to-merchant payments. The proposal simply extends that proven, co-owned model up into the merchant-services layer that PAs monopolise — and adds the one shared utility India has not yet built: a fraud-intelligence network.')

# ---------------- PSYCHOLOGY ----------------
heading('3.  The psychology of social-engineering fraud', 16)
rule_line()
body('The single most important fact: victims almost always know scams exist. Education did not fail because they were ignorant — it failed because knowledge stored yesterday is unavailable in the moment of attack. The mechanism, stage by stage:')
sub('3.1  The hijack of deliberate thought')
body('Human cognition runs on a fast, automatic, emotional system and a slow, deliberate, sceptical one. A scam’s first move is to trigger a visceral "hot state" — panic, fear, urgency or excitement. In that state, judgment narrows and the victim literally cannot access their own prior knowledge. This is the hot–cold empathy gap: the calm person who "would never fall for it" and the panicked person mid-call are, functionally, two different decision-makers.')
sub('3.2  The pressure levers')
bullet(' "I am calling from the RBI / your bank’s fraud team / the police."', bold_lead='Authority —')
bullet(' "Your account is compromised; act in ten minutes or lose everything." Digital-arrest scams weaponise this hardest.', bold_lead='Fear and urgency —')
bullet(' refunds, lotteries, jobs, investment returns.', bold_lead='Reward and greed —')
bullet(' romance scams, or a spoofed number of someone known.', bold_lead='Liking and trust —')
bullet(' a tiny first step ("just share the OTP to verify") that escalates.', bold_lead='Commitment and consistency —')
sub('3.3  Isolation and tempo')
body('The fraudster keeps the victim on the line, alone and rushed: "Don’t hang up. Don’t tell anyone — they’re in on it. Do it now." The scam is a closed loop with one voice in it; no second opinion can enter.')
sub('3.4  The funnel to a single act')
body('Every scam, whatever the script, converges on one moment: the victim authorising the payment. That is the chokepoint — everything before it is manipulation; everything after is irreversible.')
sub('The design conclusion')
body('You cannot stop people being called. You can only intervene at the moment of authorisation, and the job there is narrow and exact: break the closed loop, cool the hot state, and bring deliberate thinking back online — for two seconds, for the right person, at the right time.', italic=True)

# ---------------- AWARE LAYER ----------------
heading('4.  The Aware Layer: spot-awareness on the screen', 16)
rule_line()
body('Generic "Beware of fraud" banners fail for a precise reason: habituation. A warning that fires on every transaction becomes wallpaper within a week. The design law follows:')
body('Warnings must be rare, specific, and earned by real risk — or they are worse than nothing.', italic=True)
body('This is exactly why the Fraud Intelligence Consortium is not decoration: only shared, precise intelligence lets a warning fire only when it matters, which is the only way it keeps its power. The interventions, escalating with the consortium risk score:')

tbl = doc.add_table(rows=1, cols=2); tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
set_cell(hdr[0], 'Intervention', bold=True, white=True, size=10); shade(hdr[0], '11183A')
set_cell(hdr[1], 'What it defeats', bold=True, white=True, size=10); shade(hdr[1], '11183A')
rows = [
 ('Dynamic friction — frictionless normally; staged interruption only on high risk','Habituation; preserves UX for the 99%'),
 ('Cooling interstitial — mandatory, unskippable countdown on a high-risk new payee','The fraudster’s tempo and manufactured urgency'),
 ('Trance-breaker question — "Is someone on the phone right now telling you to pay?"','Isolation — inserts a third voice; externalises manipulation'),
 ('Named-threat warning — "This matches a digital-arrest scam. No bank asks you to move money to a safe account."','Vagueness — specificity restores salience'),
 ('Payee reputation at confirm — "This account has 9 fraud reports in 24h"','The missing second opinion at the chokepoint'),
 ('New-payee cap and ageing — small first transfer; limit rises as relationship ages','Damage from any single hijack'),
 ('Active-recall confirm — type the purpose / re-read the payee name','Autopilot tapping; re-engages deliberate thought'),
 ('Behavioural-anomaly interrupt — odd hour, rushed taps, screen-share detected','The hot state, detected from how the user transacts'),
]
for a,b in rows:
    c = tbl.add_row().cells
    set_cell(c[0], a, size=9); set_cell(c[1], b, size=9)
body('')
body('The unifying idea: the fraudster’s only power is an uninterrupted, isolated, urgent loop. Every intervention is a way for the screen to become the second voice in the room.')
body('The world is converging here, so India can leapfrog rather than pioneer alone: UK Confirmation of Payee; the UK Payment Systems Regulator’s 2024 mandatory-reimbursement rules splitting APP liability 50/50 between sending and receiving banks; Singapore’s Money Lock and Shared Responsibility Framework; Australia’s Scam-Safe Accord. India has the world’s largest real-time rail and no equivalent coordinated layer yet.')

# ---------------- FIC ----------------
heading('5.  The Fraud Intelligence Consortium (FIC)', 16)
rule_line()
body('Fraud is a network crime fought by isolated institutions. A mule account drains victims across ten banks; each bank sees only its own sliver. The criminals live in the gaps between institutions — which is precisely the function that must be made shared. The FIC is a co-owned, not-for-profit utility that does what no member can do alone:')
bullet(' every account carries a live risk score; a report at one bank reaches another bank’s confirm-screen in seconds.', bold_lead='Real-time mule / payee reputation graph —')
bullet(' continuously updated patterns that power the named-threat warnings.', bold_lead='Shared scam-typology library —')
bullet(' follows funds as they are split and hopped across banks — structurally invisible to any single institution.', bold_lead='Cross-institution layering detection —')
bullet(' banks contribute model signal without sharing raw customer data — the unlock for cooperation: privacy-safe and competition-safe.', bold_lead='Federated machine learning —')
bullet(' returns a score at the instant of authorisation; the engine of the Aware Layer.', bold_lead='Risk-score API —')
bullet(' instant cross-bank freeze and recovery of mule funds, solving today’s fatal inter-bank lag.', bold_lead='Golden-hour freeze coordination —')
bullet(' wired to the national cybercrime backbone (1930 / NCRP) for enforcement.', bold_lead='National integration —')
sub('The incentive backbone')
body('Borrow the UK design: make the receiving (mule-hosting) bank share liability for APP losses. The moment a bank pays for hosting a mule, contributing to the FIC stops being charity and becomes self-interest. The rule: a member that takes intelligence must give intelligence. That single economic design makes the consortium self-sustaining rather than a perpetually under-funded "industry good."')

# ---------------- DSC + ARCHITECTURE ----------------
heading('6.  Pillar 1: the Direct Settlement Consortium', 16)
rule_line()
body('Merchants hold a direct relationship with a consortium bank. Funds move bank-to-bank, netted over a thin switch that passes messages and computes who owes whom — but never custodies funds. That single design choice removes the entire RBI payment-aggregator escrow regime (the net-worth thresholds, dedicated escrow accounts and migration burden). Only identity (CKYC) stays central, amortised to near-zero per transaction.')
sub('What stays central vs. what distributes')
bullet(' CKYC / identity (one-time, per customer) and a thin, federated switch + net-settlement engine that never holds money.', bold_lead='Central (thin, neutral, co-owned) —')
bullet(' merchant onboarding and acquiring; fraud scoring on a bank’s own customers; direct inter-bank settlement.', bold_lead='Distributed to banks (the "pieces") —')

heading('7.  How the two pillars fit', 16)
rule_line()
body('Pillar 1 removes the rent. Pillar 2 removes the fear. The only serious objection to disintermediating PAs was "cheaper, but who handles fraud and disputes?" Pillar 2 is the answer — and it handles fraud better than the layer it replaces, because it operates at network scale a single PA never could. The Aware Layer sits on top of both: the FIC supplies the risk score, the DSC moves the money, and the screen turns that intelligence into a warning the user sees at the only moment that matters.')

# ---------------- ECONOMICS ----------------
heading('8.  Economics: who gains', 16)
rule_line()
tbl2 = doc.add_table(rows=1, cols=3); tbl2.style = 'Table Grid'
h = tbl2.rows[0].cells
for i,t in enumerate(['Stakeholder','What they gain','What they give up']):
    set_cell(h[i], t, bold=True, white=True, size=10); shade(h[i], '11183A')
econ = [
 ('Small / mid merchants','Flat near-zero cost instead of a percentage fee','—'),
 ('Banks','Acquiring + value-added revenue PAs take today; a small regulated flat fee on the rail; lower fraud losses','The comfort of blaming an intermediary'),
 ('Government','Lower UPI subsidy bill; citizen protection as political capital; payments sovereignty; forex retention','—'),
 ('Consumers','Lower prices (indirect); dramatically lower scam risk','—'),
 ('Fintechs / PAs','Repositioned, not killed — builders of the Aware Layer, the FIC’s ML, merchant SaaS','Toll-booth margins'),
]
for a,b,c in econ:
    cc = tbl2.add_row().cells
    set_cell(cc[0], a, bold=True, size=9, color=BLUE); set_cell(cc[1], b, size=9); set_cell(cc[2], c, size=9)
body('')
sub('The keystone')
body('The rail must charge a small, flat, regulated fee — not zero. Zero-MDR is precisely why banks lobby to reverse it; a flat sub-rupee fee keeps the rail self-funding without reintroducing percentage rent. This is the one number the model lives or dies on.')

# ---------------- GOVERNMENT ----------------
heading('9.  The pitch to government', 16)
rule_line()
body('Framed as three sovereign wins:')
bullet(' the State already spends roughly ₹1,500–2,000 crore a year subsidising UPI to stay free. A lower-cost, co-owned rail shrinks that bill.', bold_lead='Fiscal —')
bullet(' APP fraud is the fastest-growing financial crime and a visible public grievance. A consortium that demonstrably drives scam losses toward zero is rare, tangible political capital — protection citizens feel.', bold_lead='Citizen protection —')
bullet(' scheme fees leave the country and mule networks exploit institutional gaps. Both are fixed by infrastructure India owns and governs, extending the NPCI / UPI / CERSAI model the world already admires.', bold_lead='Strategic autonomy —')
body('The clincher: India built the hardest 80%. Bharat Direct finishes it — removing the last private toll-booth and closing the one gap UPI left open: the human at the moment of payment.', italic=True)

# ---------------- RISKS ----------------
heading('10.  Honest risks and mitigations', 16)
rule_line()
bullet(' a shared switch and FIC are systemic targets → mandate a federated, multi-operator design, not a monolith.', bold_lead='Concentration —')
bullet(' a bank consortium must not become a cartel → neutral not-for-profit governance and open membership (the NPCI precedent).', bold_lead='Antitrust / fair access —')
bullet(' sharing fraud signal must not pool raw personal data → federated learning by architecture.', bold_lead='Privacy —')
bullet(' the PA / fintech sector is jobs and tax → the repositioning narrative is essential, not cosmetic.', bold_lead='Transition cost —')
bullet(' even excellent interventions will not hit literal zero → set the target at "near-zero / world-best," backed by reimbursement so the victim is made whole regardless.', bold_lead='Residual APP fraud —')

# ---------------- CALL TO ACTION ----------------
heading('11.  Call to action: a pre-competitive working group', 16)
rule_line()
body('This is bigger than any one company — the way UPI itself began. The proposal is a pre-competitive working group:')
bullet(' RBI (Department of Payment & Settlement Systems), MeitY, the Indian Cyber Crime Coordination Centre (1930), and the Department of Financial Services.', bold_lead='Regulators / State —')
bullet(' NPCI and CERSAI.', bold_lead='Rails —')
bullet(' four to six anchor public and private banks as the consortium’s founding owners.', bold_lead='Banks —')
bullet(' the Payments Council of India and selected fintechs invited as builders of the Aware Layer and the FIC’s ML — converting potential opponents into partners.', bold_lead='Industry —')
bullet(' behavioural-science and security researchers to design and measure the interventions.', bold_lead='Academia —')
body('First deliverable: this paper plus a live pilot of the Aware Layer and a minimal FIC mule-graph on existing UPI rails — proving the fraud-reduction number before touching the settlement architecture. Lead with the fraud win; it is the easiest entry, and it earns the right to discuss disintermediation.')

# ---------------- APPENDIX ----------------
doc.add_page_break()
heading('Appendix A — Sourced figures', 15)
rule_line()
tbl3 = doc.add_table(rows=1, cols=3); tbl3.style = 'Table Grid'
h = tbl3.rows[0].cells
for i,t in enumerate(['Figure','Value','Source tier']):
    set_cell(h[i], t, bold=True, white=True, size=10); shade(h[i], '11183A')
data = [
 ('MDR on UPI & RuPay debit','Zero since 1 Jan 2020 (statutory)','Primary — law (PSS Act s.10A; IT Act s.269SU)'),
 ('Govt UPI incentive FY2024–25','₹1,500 cr at 0.15%/txn, P2M ≤₹2,000, small merchants','Primary — Cabinet (PIB)'),
 ('Govt incentive, Budget 2026–27','₹2,000 cr (industry ask ~₹10,000 cr)','Primary allocation / secondary ask'),
 ('UPI scale, Jan 2026','21.70 bn txns / ₹28.33 lakh cr in one month','Primary — NPCI'),
 ('Bank revenue forgone, zero-MDR','~₹5,500 cr/yr; govt offsets ~₹1,500–2,500 cr','Secondary — industry estimate'),
 ('PCI proposal (Mar 2025)','Reintroduce 0.3% MDR on UPI P2M > ₹2,000','Primary — industry body'),
 ('CKYC uploads','Free (Budget 2025); downloads sub-rupee, one-time','Primary'),
 ('PA escrow + net worth','₹15 cr → ₹25 cr; dedicated escrow; eff. 15 Sep 2025','Primary — RBI Master Direction'),
]
for a,b,c in data:
    cc = tbl3.add_row().cells
    set_cell(cc[0], a, size=9); set_cell(cc[1], b, size=9); set_cell(cc[2], c, size=9)
body('')
body('Note: the authoritative primary source for per-transaction cost components is the RBI Discussion Paper on Charges in Payment Systems (August 2022). For any figure marked "secondary," pin to primary RBI / NPCI data before external circulation.', italic=True, size=9)

heading('Appendix B — International precedents', 14, space_before=12)
rule_line()
bullet(' Confirmation of Payee (name-checking); Payment Systems Regulator mandatory APP reimbursement, 2024 (50/50 sending/receiving bank liability).', bold_lead='United Kingdom —')
bullet(' Money Lock; Shared Responsibility Framework for scam losses.', bold_lead='Singapore —')
bullet(' Scam-Safe Accord; mandatory Confirmation of Payee rollout.', bold_lead='Australia —')

out = r'C:\Users\dhire\OneDrive\Documents\00 Dhirendra\Innovations\Payment Innovation\Bharat Direct\Bharat-Direct-Whitepaper.docx'
doc.save(out)
print('Saved:', out)
print('Paragraphs:', len(doc.paragraphs))
