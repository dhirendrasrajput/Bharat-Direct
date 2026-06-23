# -*- coding: utf-8 -*-
"""Render Bharat Direct outreach artifacts (Concept Note, Op-ed, Manifesto) to .md and .docx."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\dhire\OneDrive\Documents\00 Dhirendra\Innovations\Payment Innovation\Bharat Direct'
NAVY=RGBColor(0x11,0x18,0x3A); BLUE=RGBColor(0x2F,0x4B,0x8F); AMBER=RGBColor(0xB0,0x6A,0x00)
GREY=RGBColor(0x55,0x5A,0x66); WHITE=RGBColor(0xFF,0xFF,0xFF)

def H1(t): return ('h1',t)
def H2(t): return ('h2',t)
def P(t):  return ('p',t)
def I(t):  return ('i',t)
def BUL(items): return ('bul',items)
def NUM(items): return ('num',items)
def TBL(h,rows): return ('table',h,rows)
def RULE(): return ('rule',)

AUTHOR = 'Dhirendra Rajput  ·  dhirendrasrajput@gmail.com  ·  June 2026'

# ============== CONCEPT NOTE (2-page spearhead) ==============
CONCEPT = [
H1('Bharat Direct — A Call to the Ecosystem'),
I('Finishing what UPI started: an aggregator-free, bank-owned payments rail and a shared Fraud Intelligence Consortium. A concept note for RBI, NPCI, iSPIRT, the banks, and anchor players.'),
RULE(),
H2('The opportunity'),
P('India built the hardest 80% of a modern payments system — UPI, a bank-owned NPCI, and central CKYC. Two inefficiencies remain. First, a private payment-aggregator layer still collects rent between merchants and banks, on rails the network no longer needs. Second, authorized push-payment (APP) / social-engineering fraud — where a person is manipulated into authorising a real payment — is the fastest-growing financial crime, and the one gap a faster rail does nothing to close.'),
H2('The proposal — two co-owned utilities'),
BUL([
 ('Direct Settlement Consortium. ','Merchants connect directly to banks; funds settle bank-to-bank over a thin, neutral switch that never custodies money; only identity (CKYC) stays central. The aggregator money-path and its escrow regime disappear. Cards and international networks are handled by a single, co-owned, at-cost gateway so merchants never hold card data (RBI tokenization already mandates this) and stay out of PCI scope.'),
 ('Fraud Intelligence Consortium (FRC). ','A shared, privacy-preserving (federated-learning) utility — a real-time mule/payee reputation graph, a scam-typology library, and golden-hour freeze coordination — that powers an on-screen "Aware Layer" delivering precise warnings at the exact moment of authorisation. Fraud is a network crime; only a shared view can see it whole.'),
]),
H2('Why now, why India'),
P('India has the world’s largest real-time rail, a national cybercrime backbone (1930 / NCRP / I4C), and telecom-fraud tooling — but no utility fusing them. The world is converging on these ideas (UK mandatory APP reimbursement, Singapore’s Anti-Scam Centre and Shared Responsibility Framework, Australia’s AFCX). India can build the most complete version, on infrastructure it already owns and governs.'),
H2('What is being built today vs. what needs the ecosystem'),
BUL([
 ('Buildable now (in the open). ','An interoperable, drop-in fraud-intelligence layer and an open protocol — designed so contributors and consumers integrate without changing their core systems. This proves the fraud-reduction number on existing rails, with one bank or PSP, needing no one’s permission.'),
 ('Needs the ecosystem. ','Native split-settlement (an NPCI rail feature), a production-grade consortium switch (an owned, open-source alternative to licensed monoliths, operated at cost), and a real FRC with live bank data-sharing (a governance and legal change).'),
]),
H2('The ask'),
NUM([
 'A pre-competitive working group — convened by iSPIRT or NPCI, with RBI (DPSS), DFS, I4C, anchor banks, and the payments industry.',
 'A regulatory sandbox / RBIH pilot of the Aware Layer + a minimal mule-graph on existing UPI rails — proving the fraud win first.',
 'An anchor owner (a bank consortium, or a player such as Reliance or Tata) to carry the rail at national scale.',
]),
P('This is a flag, not a pitch for sale. The blueprint is meant to become shared public infrastructure. The author can architect it and is building the fraud layer in the open; owning national infrastructure needs the institutions named above.'),
H2('Key facts (sourced)'),
TBL(['Figure','Value','Source'],[
 ['MDR on UPI / RuPay debit','Zero since 1 Jan 2020 (statutory)','PSS Act s.10A; IT Act s.269SU'],
 ['Govt UPI incentive FY2024-25','Rs 1,500 cr at 0.15%/txn, P2M <= Rs 2,000','PIB (Cabinet)'],
 ['UPI scale, Jan 2026','21.70 bn txns / Rs 28.33 lakh cr in one month','NPCI'],
 ['Per-txn cost reference','RBI Discussion Paper on Charges in Payment Systems, Aug 2022','RBI'],
]),
RULE(),
I('Author: ' + AUTHOR + '. Full white paper, deck, and open protocol specification available on request.'),
]

# ============== OP-ED (~600 words) ==============
OPED = [
H1('India built UPI. It’s time to finish the job.'),
I('Op-ed draft (~600 words). Byline: ' + AUTHOR),
RULE(),
P('A decade ago, India did something no other country had managed: it made moving money as free and instant as sending a text. UPI now carries more than 21 billion transactions a month. We rightly celebrate it. But two costs are still hiding in plain sight on every payment we make — and both are now fixable.'),
P('The first is a toll-booth. When you pay a merchant, a private intermediary — a payment aggregator — often still takes a slice, even though the money itself can move directly from your bank to the merchant’s. For the smallest shopkeeper this is a real tax; for the largest merchant it is a layer of cost and fragility the network outgrew years ago. UPI already proves the money can move without a middleman in the path. We simply never finished extending that logic to the services built on top.'),
P('The second cost is darker, and it is growing. The fastest-rising financial crime in India is not someone hacking your account. It is someone talking you into paying them. A caller claiming to be from your bank, the police, or a courier company; a fake job, a fake refund, a fake arrest. These scams do not break the system — they break the person, in a moment of panic, into authorising a payment themselves. A faster rail does nothing to stop this. You cannot out-engineer a phone call.'),
P('Here is why this matters now. Both problems share a root cause: things that should be shared are still siloed. Fraud is a network crime — a single mule account drains victims across a dozen banks — but each bank sees only its own sliver, while the criminal sees the whole map. And the merchant-services layer that should be a neutral public utility is instead a patchwork of private toll-booths.'),
P('The answer is not a new app or another regulator circular. It is two pieces of shared infrastructure, built the way UPI itself was built — co-owned by banks, neutral, and not-for-profit.'),
P('The first is a direct settlement rail: merchants connect to banks, money settles bank-to-bank, and only identity stays central. International cards keep working through a single at-cost gateway, so no merchant ever holds your card data. The private margin disappears; the function stays where it is genuinely needed.'),
P('The second is a Fraud Intelligence Consortium — a shared, privacy-preserving early-warning system. When a mule account is flagged at one bank, every bank knows within seconds. And at the one moment that matters — when you are about to authorise a payment — your screen becomes the second voice in the room: “This account has nine fraud reports today,” or “This looks like a digital-arrest scam; no real bank asks you to move money to a safe account.” Not another ignored banner, but a precise, rare warning, fired only when shared intelligence says the risk is real.'),
P('None of this requires importing foreign technology or starting from scratch. The world is already moving this way — Britain now splits fraud liability between sending and receiving banks; Singapore co-locates banks and police in an anti-scam command. India, with the largest real-time rail on earth and a national cybercrime helpline already in place, is better positioned than any of them to build the most complete version.'),
P('What it needs is will, and the right people in one room: the RBI, NPCI, the banks, the cybercrime agencies, and the industry — the same pre-competitive coalition that gave us UPI. Start with the fraud win, because it protects citizens and no one can argue against it. Let it earn the right to finish the rest.'),
P('India taught the world how to move money. The next lesson is how to protect the people moving it — and how to stop charging them for plumbing they already own.'),
]

# ============== MANIFESTO (repo-native md) ==============
MANIFESTO = [
H1('Bharat Direct'),
I('An open blueprint to finish what UPI started — and an open fraud-intelligence layer being built in the open, right now.'),
RULE(),
H2('The problem'),
P('India built UPI, NPCI, and CKYC — the hardest 80%. Two things remain. A private aggregator layer still taxes payments on rails the network no longer needs. And authorized push-payment (APP) fraud — scams that manipulate the person, not the system — is the fastest-growing financial crime, untouched by any faster rail.'),
H2('The blueprint'),
BUL([
 ('Direct Settlement Consortium. ','Bank-to-bank merchant payments, a thin neutral switch, no aggregator in the money path, cards handled by one at-cost gateway so no merchant holds card data.'),
 ('Fraud Intelligence Consortium. ','A shared, privacy-preserving risk graph + an on-screen Aware Layer that warns users at the moment of authorisation — rare, specific, and earned by real risk.'),
]),
H2('What is open here'),
BUL([
 ('The protocol. ','An open spec for direct pay-by-bank, risk-aware authorisation, and programmable split settlement. See SPEC.md.'),
 ('The reference fraud layer. ','A drop-in, federated, interoperable fraud-intelligence service + Aware Layer SDK.'),
 ('The integration promise. ','Contributors and consumers integrate without ripping out their core systems — one API, conformance profiles, federated by design, easy in and easy out.'),
]),
H2('Who should build the rest'),
P('Owning national infrastructure is not a solo job. The rail and the full consortium need NPCI, the banks, RBI’s blessing, or an anchor with national scale. This repository is the blueprint and the fraud wedge — a flag for those who can carry it the rest of the way.'),
H2('How to join'),
NUM([
 'Read SPEC.md and the white paper.',
 'Build against the mock FRC and the Aware Layer SDK — prove the fraud win on your own rails.',
 'Open an issue or reach out as a design partner: a bank, a PSP, a fraud team, or an anchor owner.',
]),
RULE(),
I('This is meant to become shared public infrastructure. Contributions and critique are the point. Spec: CC BY 4.0. Code: Apache-2.0. Author: ' + AUTHOR + '.'),
]

# ---------- renderers ----------
def to_md(blocks):
    out=[]
    for b in blocks:
        k=b[0]
        if k=='h1': out.append(f'# {b[1]}\n')
        elif k=='h2': out.append(f'## {b[1]}\n')
        elif k=='p': out.append(f'{b[1]}\n')
        elif k=='i': out.append(f'*{b[1]}*\n')
        elif k=='rule': out.append('---\n')
        elif k=='bul':
            for it in b[1]:
                if isinstance(it,tuple): out.append(f'- **{it[0].rstrip()}** {it[1].lstrip()}')
                else: out.append(f'- {it}')
            out.append('')
        elif k=='num':
            for i,it in enumerate(b[1]): out.append(f'{i+1}. {it}')
            out.append('')
        elif k=='table':
            out.append('| '+' | '.join(b[1])+' |')
            out.append('| '+' | '.join(['---']*len(b[1]))+' |')
            for r in b[2]: out.append('| '+' | '.join(r)+' |')
            out.append('')
    return '\n'.join(out)+'\n'

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def to_docx(blocks, path):
    doc=Document(); n=doc.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(11)
    n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.15
    for b in blocks:
        k=b[0]
        if k=='h1':
            p=doc.add_paragraph(); r=p.add_run(b[1]); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
            p.paragraph_format.space_after=Pt(8)
        elif k=='h2':
            p=doc.add_paragraph(); r=p.add_run(b[1]); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=BLUE
            p.paragraph_format.space_before=Pt(12)
        elif k=='p': doc.add_paragraph().add_run(b[1]).font.size=Pt(11)
        elif k=='i':
            p=doc.add_paragraph(); r=p.add_run(b[1]); r.italic=True; r.font.size=Pt(10.5); r.font.color.rgb=GREY
        elif k=='rule':
            p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr(); pb=OxmlElement('w:pBdr')
            bo=OxmlElement('w:bottom'); bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'6')
            bo.set(qn('w:space'),'1'); bo.set(qn('w:color'),'2F4B8F'); pb.append(bo); pPr.append(pb)
        elif k=='bul':
            for it in b[1]:
                p=doc.add_paragraph(style='List Bullet')
                if isinstance(it,tuple):
                    r=p.add_run(it[0]); r.bold=True; p.add_run(it[1])
                else: p.add_run(it)
        elif k=='num':
            for it in b[1]: doc.add_paragraph(it, style='List Number')
        elif k=='table':
            hdr=b[1]; rows=b[2]; t=doc.add_table(rows=1,cols=len(hdr)); t.style='Table Grid'
            for i,htext in enumerate(hdr):
                c=t.rows[0].cells[i]; shade(c,'11183A'); c.text=''
                rr=c.paragraphs[0].add_run(htext); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=WHITE
            for row in rows:
                cells=t.add_row().cells
                for i,val in enumerate(row):
                    cells[i].text=''; cells[i].paragraphs[0].add_run(val).font.size=Pt(9.5)
            doc.add_paragraph()
    doc.save(path)

# ---------- write ----------
with open(os.path.join(OUT,'Concept-Note.md'),'w',encoding='utf-8') as f: f.write(to_md(CONCEPT))
with open(os.path.join(OUT,'Op-Ed.md'),'w',encoding='utf-8') as f: f.write(to_md(OPED))
with open(os.path.join(OUT,'MANIFESTO.md'),'w',encoding='utf-8') as f: f.write(to_md(MANIFESTO))
to_docx(CONCEPT, os.path.join(OUT,'Bharat-Direct-Concept-Note.docx'))
to_docx(OPED, os.path.join(OUT,'Bharat-Direct-OpEd.docx'))

wc=len(' '.join(b[1] for b in OPED if b[0] in ('p',)).split())
print('Wrote Concept-Note (.md/.docx), Op-Ed (.md/.docx), MANIFESTO.md')
print('Op-ed body word count:', wc)
